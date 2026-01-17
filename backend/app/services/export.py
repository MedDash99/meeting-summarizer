from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

from app.models.schemas import MeetingSummary


def create_word_document(summary: MeetingSummary) -> str:
    """Generate a formatted Word document from meeting summary."""
    doc = Document()
    
    # Title
    title = doc.add_heading(summary.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(summary.summary)
    
    # Participants (only if present)
    if summary.participants:
        doc.add_heading("Participants", level=1)
        for p in summary.participants:
            name_str = p.name or "Unknown Speaker"
            role_str = f" ({p.role})" if p.role else ""
            doc.add_paragraph(f"• {name_str}{role_str}")
    
    # Key Points (only if present)
    if summary.key_points:
        doc.add_heading("Key Discussion Points", level=1)
        for point in summary.key_points:
            doc.add_paragraph(f"• {point}")
    
    # Decisions
    if summary.decisions:
        doc.add_heading("Decisions Made", level=1)
        for d in summary.decisions:
            p = doc.add_paragraph(f"✓ {d.description}")
            if d.context:
                context_para = doc.add_paragraph(f"   Context: {d.context}")
                for run in context_para.runs:
                    run.italic = True
    
    # Action Items
    if summary.action_items:
        doc.add_heading("Action Items", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = "Task"
        header_cells[1].text = "Assignee"
        header_cells[2].text = "Deadline"
        
        for item in summary.action_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.task
            row_cells[1].text = item.assignee or "TBD"
            row_cells[2].text = item.deadline or "TBD"
    
    # Full Transcript
    doc.add_page_break()
    doc.add_heading("Full Transcript", level=1)
    doc.add_paragraph(summary.transcript)
    
    # Save to temp file
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name
